# -*- coding: utf-8 -*-
import binascii
import glob
from os import getcwd, listdir
from os.path import isfile, join
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

smo_map = {161: '\xe0\xb8\x81', 162: '\xe0\xb8\x82', 163: '\xe0\xb8\x84', 164: '\xe0\xb8\x86', 165: '\xe0\xb8\x87', 166: '\xe0\xb8\x88', 167: '\xe0\xb8\x89', 168: '\xe0\xb8\x8a', 169: '\xe0\xb8\x8b', 170: '\xe0\xb8\x8c', 171: '\xe0\xb8\x8d', 172: '\xe0\xb8\x8e', 173: '\xe0\xb8\x8f', 174: '\xe0\xb8\x90', 175: '\xe0\xb8\x91', 176: '\xe0\xb8\x92', 177: '\xe0\xb8\x93', 178: '\xe0\xb8\x94', 179: '\xe0\xb8\x95', 180: '\xe0\xb8\x96', 181: '\xe0\xb8\x97', 182: '\xe0\xb8\x98', 183: '\xe0\xb8\x99', 184: '\xe0\xb8\x9a', 185: '\xe0\xb8\x9b', 186: '\xe0\xb8\x9c', 187: '\xe0\xb8\x9d', 188: '\xe0\xb8\x9e', 189: '\xe0\xb8\x9f', 190: '\xe0\xb8\xa0', 191: '\xe0\xb8\xa1', 192: '\xe0\xb8\xa2', 193: '\xe0\xb8\xa3', 194: '\xe0\xb8\xa4', 195: '\xe0\xb8\xa5', 196: '\xe0\xb8\xa7', 197: '\xe0\xb8\xa8', 198: '\xe0\xb8\xa9', 199: '\xe0\xb8\xaa', 200: '\xe0\xb8\xab', 201: '\xe0\xb8\xac', 202: '\xe0\xb8\xad', 203: '\xe0\xb8\xae'}
smo_map.update({144:'๐', 145:'๑', 146:'๒', 147:'๓', 148:'๔', 149:'๕', 150:'๖', 151:'๗', 152:'๘', 153:'๙'})
smo_map.update({204:'ะ', 206:'า', 217:'ิ', 218:'ี', 210:'โ', 219:'ึ', 220:'ื', 215:'ุ', 216:'ู', 222:'ํ', 228:'์', 223:'็'})
smo_map.update({224:'่', 225:'้', 226:'๊', 227:'๋', 229:'ํ', 207:'ำ', 221:'ั', 208:'เ', 209:'แ', 213:'ๆ', 214:'ฯ', 211:'ใ', 212:'ไ'})
stateDict = {'\\h':'heading', '\\mt':'sub-heading', '\\c':'chapter', '\\p\\v':'paragraph', '\\id':'id', '\\p':'paragraph', '\\v':'paragraph', '\\s':'sub-heading'}


def deleteTags(string):
	tags = ['\\v', 'n', '\\tx', '\\pn', '\\p', '\\fe']
	for tag in tags:
		string = string.replace(tag, '')
	return string

def readFile(filePath):
	f = open(filePath, 'rb')
	txtBuffer = ""
	fullText = []
	state = ""
	try:
		byte = int(binascii.hexlify(f.read(1)), 16)
		while byte != "":
			if byte in smo_map:
				txtBuffer += smo_map[byte]
			elif byte >= 32 and byte <= 126 and not(chr(byte) == '<' or chr(byte) == '>'):
				txtBuffer += chr(byte)
			elif byte == 13:
				txtBuffer += ' '
			if len(txtBuffer) > 1 and '\\' in txtBuffer and txtBuffer[-1] == ' ':
				if txtBuffer[txtBuffer.find('\\'):-1] in stateDict:
					if state != '':
						fullText += [(state, txtBuffer[:txtBuffer.find('\\')])] if txtBuffer[:txtBuffer.find('\\')].isspace else []
					state = stateDict[txtBuffer[txtBuffer.find('\\'):-1]]
					txtBuffer = ''
				else:
					txtBuffer = txtBuffer[:txtBuffer.find('\\')]
			elif txtBuffer[-3:] == '\\fe':
				break
			byte = int(binascii.hexlify(f.read(1)), 16)
	except:
		print "[!] Got a little error with this file \'%s\'." % filePath
	finally:
		fullText += [(state, txtBuffer[:txtBuffer.find('\\')])]
		f.close()
	return fullText


def convert(name, fullText):
	document = Document()

	for cType, content in fullText:
		if content.isspace():
			continue
		elif cType == 'heading':
			p = document.add_paragraph()
			p = p.add_run(unicode(content, 'utf-8')).font
			p.size = Pt(24)
			p.name = 'LawafontUPC'
			p.bold = True
		elif cType == 'paragraph':
			p = document.add_paragraph()
			p = p.add_run(unicode(content, 'utf-8')).font
			p.name = 'LawafontUPC'
			p.size = Pt(18)
		elif cType == 'chapter':
			p = document.add_paragraph()
			p_format = p.paragraph_format
			p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
			p = p.add_run(unicode('บทที่ ' + content, 'utf-8')).font
			p.name = 'LawafontUPC'
			p.size = Pt(24)
			p.bold = True
		elif cType == 'sub-heading':
			p = document.add_paragraph()
			p_format = p.paragraph_format
			p_format.space_before = Pt(28)
			p = p.add_run(unicode(content, 'utf-8')).font
			p.name = 'LawafontUPC'
			p.size = Pt(18)
			p.bold = True

	document.save(name.split('.')[0]+'.docx')

def execute(path):
	print '\n[+] Current Directory is \'%s\'' % path
	directories = [join(path, i) for i in listdir(path) if not isfile(join(path, i))]
	texFiles = glob.glob(path + '/*.TEX')
	for fileName in texFiles:
		fullText = readFile(fileName)
		t = fileName.split('/')[-1]
		convert(fileName, fullText)
		print '[+] File \'%s\' is converted to \'%s\' already!' % (fileName.split('/')[-1], t.split('.')[0] + '.docx')
	for d in directories:
		execute(d)

execute(getcwd())